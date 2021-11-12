import pandas
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from datetime import datetime


def change_orientation():
    current_section = word_document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = word_document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height

    return new_section


def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


tables = pandas.read_excel(r"Header and metadata list (10).xlsx", sheet_name=None)

table_names_to_export = list(tables.keys())[3:]

word_document = Document()
document_name = 'Header and Metadata'

word_document.add_heading(document_name, 0)

p = word_document.add_paragraph('')
p.add_run(f"Generation timestamp: {datetime.utcnow()}Z").italic = True

section = change_orientation()
margin = 0.5
section.top_margin = Cm(margin)
section.bottom_margin = Cm(margin)
section.left_margin = Cm(margin)
section.right_margin = Cm(margin)

for table_name in table_names_to_export:

    # Add heading for table
    word_document.add_heading(table_name, level=1)



    data_table = tables[table_name]

    #print(data_table.columns)

    if not "Unnamed" in data_table.columns[2]:
        explanation_text = f"Header metadata fields for {data_table.columns[2]} sent by {data_table.columns[4]}"
        print(explanation_text)
        word_document.add_paragraph(explanation_text)

    # customizing the table
    table = word_document.add_table(0, 0)  # we add rows iteratively
    table.style = 'TableGrid'
    # table.set_repeat_table_header(table.rows[0])

    data_table.columns = data_table.iloc[1]
    data_table = data_table[2:]

    if "Example" in data_table.columns:
        data_table = data_table.drop(["Example"], axis=1)

    data_table = data_table.dropna(axis=1)
    data_table["Comments"] = ""

    table.add_row()
    for column_number, column_name in enumerate(data_table.columns):
        table.add_column(Cm(5))
        table.cell(0, column_number).text = column_name

    make_rows_bold(table.rows[0])

    for row_number, row_data in enumerate(data_table.itertuples()):

        row = table.add_row()

        for cell_index, value in enumerate(row_data):

            row.cells[cell_index-1].text = str(value)

    word_document.add_page_break()

word_document.save(document_name + '.docx')